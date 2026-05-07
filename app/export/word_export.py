# app/export/word_export.py
"""
Word Export v4 — CLEAN DOCX GENERATION

Rules:
- heading → doc.add_heading(text, level=2)
- list    → doc.add_paragraph(text, style='List Number')
- text    → doc.add_paragraph(text) with justify + first-line indent
- table   → doc.add_table()

NO duplicate list markers.
NO noise in output.
Tight margins: 1.5cm left, 1cm right.
"""

import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from typing import List, Dict

FONT_NAME = "Times New Roman"


# ============================================================
# NOISE FILTER (final safety net)
# ============================================================

_NOISE_RE = [
    re.compile(r'^\d{1,2}:\d{2}\s*[\(\)\d♡⊙☆\s]*$'),
    re.compile(r'^[45]G[\+]?\s*'),
    re.compile(r'KB/[sS]', re.IGNORECASE),
    re.compile(r'^\s*[\(\s]*\d{1,3}\s*%'),
    re.compile(r'^\s*\d{1,3}\s*$'),
    re.compile(r'^\s*<?\s*\d{1,4}/\d{1,4}'),
    re.compile(r'Scan[_\s]?\d{8}'),
    re.compile(r'chrome-native://'),
    re.compile(r'OVERFLOWED'),
]


def _is_noise(text: str) -> bool:
    t = text.strip()
    if not t or len(t) <= 1:
        return True
    if len(t) <= 3 and re.match(r'^[\d\s.,;:()\-]+$', t):
        return True
    for p in _NOISE_RE:
        if p.search(t):
            return True
    return False


# ============================================================
# FONT
# ============================================================

def _set_font(run, name=FONT_NAME, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)


def _setup_doc(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1)


# ============================================================
# BLOCK WRITERS
# ============================================================

def _add_heading(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_font(run, size=16 if level == 1 else 14, bold=True)
    # Headings should never look like indented body paragraphs
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.first_line_indent = Cm(0)
    h.paragraph_format.left_indent = Cm(0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.line_spacing = Pt(18)


def _add_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    # Consistent body spacing (avoid cramped vs overly spaced)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    # Reduce indent; 1cm was too aggressive for many mobile OCR outputs
    p.paragraph_format.first_line_indent = Cm(0.6)
    p.paragraph_format.left_indent = Cm(0)


def _add_list_item(doc, text, *, kind: str | None = None):
    """
    Add a single list item. Text should be CLEAN (no "1." prefix).
    The prefix is handled by Word's List Number style.
    """
    # Double-check: remove any remaining number prefix
    clean = re.sub(r'^\d{1,3}[\.\)\-]\s*', '', text)
    clean = re.sub(r'^[\-•·*‣▪►]\s*', '', clean)
    clean = re.sub(r'^[a-zA-Z][\.\)]\s*', '', clean)
    clean = clean.strip()

    if not clean:
        return

    style = 'List Bullet' if kind in ("list_bullet", "bullet") else 'List Number'
    p = doc.add_paragraph(style=style)
    run = p.add_run(clean)
    _set_font(run, size=12)
    # Keep list spacing consistent with paragraphs
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    # Ensure wrapped lines (continuations) align properly under the text,
    # not under the marker/number.
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_table(doc, table_data):
    if not table_data or not table_data[0]:
        return
    rows = len(table_data)
    cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = str(cell_text) if cell_text else ""
                for p in cell.paragraphs:
                    for run in p.runs:
                        _set_font(run, size=10, bold=(i == 0))
    doc.add_paragraph()


# ============================================================
# MAIN BUILDERS
# ============================================================

def build_docx_from_blocks(blocks: List[Dict]) -> bytes:
    """
    Build DOCX from structured blocks.
    Each block: {"type": "heading"|"list"|"text"|"table", "content": str|list}
    """
    doc = Document()
    _setup_doc(doc)

    for block in blocks:
        btype = block.get("type", "text")
        content = block.get("content", "")
        if not content:
            continue

        content_str = str(content)

        # Final noise filter
        if btype not in ("table", "figure") and _is_noise(content_str):
            continue

        if btype == "title":
            _add_heading(doc, content_str, level=1)
        elif btype == "heading":
            _add_heading(doc, content_str, level=2)
        elif btype == "list":
            _add_list_item(doc, content_str, kind=block.get("list_kind"))
        elif btype == "table":
            if isinstance(content, list):
                _add_table(doc, content)
            else:
                _add_paragraph(doc, content_str)
        elif btype != "figure":
            _add_paragraph(doc, content_str)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx_from_text(text: str) -> bytes:
    doc = Document()
    _setup_doc(doc)
    for line in text.split("\n"):
        line = line.strip()
        if not line or _is_noise(line):
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        _set_font(run, size=12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_multi_page_docx(pages: List[List[Dict]]) -> bytes:
    doc = Document()
    _setup_doc(doc)
    for idx, blocks in enumerate(pages):
        if idx > 0:
            doc.add_page_break()
        for block in blocks:
            btype = block.get("type", "text")
            content = block.get("content", "")
            if not content:
                continue
            content_str = str(content)
            if btype not in ("table", "figure") and _is_noise(content_str):
                continue
            if btype == "title":
                _add_heading(doc, content_str, level=1)
            elif btype == "heading":
                _add_heading(doc, content_str, level=2)
            elif btype == "list":
                _add_list_item(doc, content_str, kind=block.get("list_kind"))
            elif btype == "table":
                if isinstance(content, list):
                    _add_table(doc, content)
                else:
                    _add_paragraph(doc, content_str)
            elif btype != "figure":
                _add_paragraph(doc, content_str)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()