# app/ocr_service.py
# FIXED + PADDLEOCR VERSION

import cv2
import numpy as np
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile
import time
from io import BytesIO

FONT_NAME = "Times New Roman"

# ============================================================
# INIT OCR (GLOBAL)
# ============================================================

ocr = PaddleOCR(
    use_angle_cls=True,
    lang='ru'   # Uzbek + Russian uchun eng yaxshi
)

# ============================================================
# DOCUMENT DETECTION
# ============================================================

def detect_document(image):

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    blur = cv2.GaussianBlur(gray, (5,5), 0)

    edges = cv2.Canny(blur, 75, 200)

    contours, _ = cv2.findContours(
        edges,
        cv2.RETR_LIST,
        cv2.CHAIN_APPROX_SIMPLE
    )

    contours = sorted(
        contours,
        key=cv2.contourArea,
        reverse=True
    )

    for c in contours[:5]:

        peri = cv2.arcLength(c, True)

        approx = cv2.approxPolyDP(
            c,
            0.02 * peri,
            True
        )

        if len(approx) == 4:

            pts = approx.reshape(4,2)

            return four_point_transform(image, pts)

    return image

# ============================================================

def four_point_transform(image, pts):

    rect = order_points(pts)

    (tl, tr, br, bl) = rect

    widthA = np.linalg.norm(br-bl)
    widthB = np.linalg.norm(tr-tl)

    maxWidth = max(int(widthA), int(widthB))

    heightA = np.linalg.norm(tr-br)
    heightB = np.linalg.norm(tl-bl)

    maxHeight = max(int(heightA), int(heightB))

    dst = np.array([
        [0,0],
        [maxWidth-1,0],
        [maxWidth-1,maxHeight-1],
        [0,maxHeight-1]
    ], dtype="float32")

    M = cv2.getPerspectiveTransform(rect, dst)

    warped = cv2.warpPerspective(
        image,
        M,
        (maxWidth,maxHeight)
    )

    return warped

# ============================================================

def order_points(pts):

    rect = np.zeros((4,2), dtype="float32")

    s = pts.sum(axis=1)

    rect[0] = pts[np.argmin(s)]
    rect[2] = pts[np.argmax(s)]

    diff = np.diff(pts, axis=1)

    rect[1] = pts[np.argmin(diff)]
    rect[3] = pts[np.argmax(diff)]

    return rect

# ============================================================
# PREPROCESS
# ============================================================

def preprocess_image(image):

    image = detect_document(image)

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    h, w = gray.shape

    if max(h, w) < 3000:

        scale = 3000 / max(h, w)

        gray = cv2.resize(
            gray,
            None,
            fx=scale,
            fy=scale,
            interpolation=cv2.INTER_CUBIC
        )

    clahe = cv2.createCLAHE(
        clipLimit=2.0,
        tileGridSize=(8,8)
    )

    gray = clahe.apply(gray)

    thresh = cv2.adaptiveThreshold(
        gray,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        10
    )

    return thresh

# ============================================================
# OCR
# ============================================================

def run_ocr(image_path):

    image = cv2.imread(image_path)

    result = ocr.ocr(image, cls=True)

    lines = []

    for line in result:

        for word in line:

            text = word[1][0]

            if text.strip():

                lines.append(text)

    final_text = "\n".join(lines)

    return final_text.strip()

# ============================================================
# OCR + CONFIDENCE
# ============================================================

def ocr_with_stats(image_path):

    start = time.time()

    image = cv2.imread(image_path)

    h, w = image.shape[:2]

    processed = preprocess_image(image)

    result = ocr.ocr(processed, cls=True)

    lines = []

    confidences = []

    for line in result:

        for word in line:

            text = word[1][0]

            conf = word[1][1]

            if text.strip():

                lines.append(text)

                confidences.append(conf)

    text = "\n".join(lines)

    avg_conf = (
        sum(confidences) / len(confidences)
        if confidences else 0
    )

    elapsed = time.time() - start

    return {
        "text": text,
        "confidence": round(avg_conf, 1),
        "processing_time": round(elapsed, 2),
        "image_size": {
            "width": w,
            "height": h
        },
    }

# ============================================================
# DOCX
# ============================================================

def create_docx(text, output_path=None):

    doc = Document()

    for section in doc.sections:

        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    for line in text.split("\n"):

        if not line.strip():
            continue

        p = doc.add_paragraph()

        run = p.add_run(line)

        run.font.size = Pt(12)

        run.font.name = FONT_NAME

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if output_path:

        doc.save(output_path)

        return output_path

    buf = BytesIO()

    doc.save(buf)

    return buf.getvalue()