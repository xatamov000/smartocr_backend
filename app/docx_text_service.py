from io import BytesIO
from docx import Document
from docx.shared import Pt

# Times New Roman o‘rnini bosuvchi (Linux serverlar uchun)
FONT_NAME = "Liberation Serif"


def _apply_font(run, *, size: int | None = None, bold: bool = False, italic: bool = False):
    """
    ✅ SAFE font apply (Render / Docker friendly)
    - Cyrillic / Uzbek / Russian OK
    - Bold / Italic OK
    - No XML hacks
    """
    run.font.name = FONT_NAME
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _is_heading(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    return len(s) <= 60 and sum(1 for c in s if c.isupper()) >= max(3, int(len(s) * 0.25))


def _is_numbered(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if s[0].isdigit():
        if len(s) >= 2 and s[1] in [".", ")", " "]:
            return True
        if len(s) >= 3 and s[1].isdigit() and s[2] in [".", ")", " "]:
            return True
    return False


def _is_bullet(line: str) -> bool:
    return line.strip().startswith(("-", "•", "·", "*"))


def build_docx_bytes_from_text(text: str) -> bytes:
    """
    Text -> DOCX bytes
    ✅ No 500 error
    ✅ Cyrillic safe
    ✅ Bold / Italic works
    """
    doc = Document()
    text = (text or "").replace("\r\n", "\n").strip("\n")

    if not text.strip():
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    lines = text.split("\n")
    blank = False

    for raw in lines:
        line = raw.rstrip()

        if not line.strip():
            blank = True
            continue

        if blank:
            doc.add_paragraph("")
            blank = False

        s = line.strip()

        if _is_numbered(s):
            p = doc.add_paragraph("", style="List Number")
        elif _is_bullet(s):
            p = doc.add_paragraph("", style="List Bullet")
        elif _is_heading(s):
            p = doc.add_paragraph("")
        else:
            p = doc.add_paragraph("")

        if _is_heading(s):
            run = p.add_run(s)
            _apply_font(run, size=14, bold=True)
        else:
            run = p.add_run(s)
            _apply_font(run, size=12)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
