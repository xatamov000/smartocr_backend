"""
🔥 YANGILANGAN DOCX SERVICE - FORMAT SAQLASH BILAN

Bu service pytesseract'ning image_to_data() funksiyasidan foydalanib:
1. Har bir so'zning koordinatalari va o'lchamlarini oladi
2. Shriftlar, abzatzlar, sarlavhalarni aniqlaydi
3. Formatlashni saqlagan holda DOCX yaratadi
"""

from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path
from typing import List, Dict, Tuple, Optional

import pytesseract
from pytesseract import Output
from PIL import Image, ImageOps, ImageFilter

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============================================================
# Config
# ============================================================

FONT_NAME = "Arial"  # Universal font
DEFAULT_FONT_SIZE = 11
MIN_FONT_SIZE = 9
MAX_FONT_SIZE = 28

# ============================================================
# Tesseract setup
# ============================================================

try:
    from .ocr_service_improved import ensure_tesseract
except Exception:
    from ocr_service_improved import ensure_tesseract


# ============================================================
# OCR helpers
# ============================================================

def _tess_config(psm: int = 3) -> str:
    """PSM 3 = automatic page segmentation"""
    return f"--oem 3 --psm {psm} --dpi 300"


def _maybe_autorotate(img: Image.Image) -> Image.Image:
    try:
        osd = pytesseract.image_to_osd(img, output_type=Output.DICT)
        rotate = int(osd.get("rotate", 0) or 0)
        if rotate in (90, 180, 270):
            return img.rotate(360 - rotate, expand=True)
    except Exception:
        pass
    return img


def _preprocess_for_ocr(img: Image.Image, fast_mode: bool = False) -> Image.Image:
    """Rasmni OCR uchun tayyorlash"""
    img = ImageOps.exif_transpose(img)

    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    if img.mode != "L":
        img = img.convert("L")

    w, h = img.size
    max_dim = max(w, h)
    
    if max_dim < 1000:
        img = img.resize((w * 2, h * 2), Image.Resampling.LANCZOS)
    elif max_dim < 1400 and not fast_mode:
        scale = 1.5
        img = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)
    elif max_dim > 3000:
        scale = 2000 / max_dim
        img = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)

    if not fast_mode:
        img = ImageOps.autocontrast(img)
        img = img.filter(ImageFilter.MedianFilter(size=3))

    img = img.point(lambda p: 255 if p > 160 else 0)
    return img


def _safe_open_image(path: Path) -> Image.Image:
    img = Image.open(path)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    return img


# ============================================================
# Font & Style Utils
# ============================================================

def _px_to_pt(px: float, dpi: int = 300) -> float:
    """Pixel dan Point ga o'tkazish (1pt = 1/72 inch)"""
    inches = px / dpi
    pt = inches * 72
    return max(MIN_FONT_SIZE, min(MAX_FONT_SIZE, pt))


def _apply_font(run, font_name: str = FONT_NAME):
    """Universal font qo'llash (Cyrillic safe)"""
    run.font.name = font_name
    
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _is_heading_text(text: str, height_ratio: float = 1.0) -> bool:
    """Sarlavha ekanligini aniqlash"""
    text = text.strip()
    if not text or len(text) > 100:
        return False
    
    # Agar juda katta shrift bo'lsa
    if height_ratio >= 1.5:
        return True
    
    # Agar qisqa va katta harflar ko'p bo'lsa
    if len(text) <= 60:
        upper_count = sum(1 for c in text if c.isupper())
        if upper_count >= max(3, int(len(text) * 0.3)):
            return True
    
    return False


def _is_bullet_or_numbered(text: str) -> Tuple[bool, str]:
    """List item ekanligini aniqlash"""
    text = text.strip()
    
    # Bullet list
    if text.startswith(("•", "-", "·", "*", "○", "■", "►")):
        return True, "bullet"
    
    # Numbered list
    if len(text) >= 2 and text[0].isdigit():
        if text[1] in (".", ")", " "):
            return True, "number"
        if len(text) >= 3 and text[1].isdigit() and text[2] in (".", ")", " "):
            return True, "number"
    
    return False, ""


# ============================================================
# Extract structured data from image
# ============================================================

def _extract_lines_with_format(img: Image.Image, lang: str) -> List[Dict]:
    """
    Pytesseract image_to_data bilan har bir so'z va satrni oladi.
    Har bir satr uchun: text, left, top, width, height, format
    """
    ensure_tesseract()
    
    img = _maybe_autorotate(_preprocess_for_ocr(img))
    
    # image_to_data - har bir so'z uchun ma'lumot
    data = pytesseract.image_to_data(
        img,
        lang=lang,
        output_type=Output.DICT,
        config=_tess_config(),
    )
    
    # So'zlarni satrlarga birlashtirish
    lines_dict = {}  # key: (block, par, line)
    
    for i, txt in enumerate(data["text"]):
        txt = (txt or "").strip()
        try:
            conf = int(float(data["conf"][i]))
        except Exception:
            conf = -1
        
        if not txt or conf < 0:
            continue
        
        block = data["block_num"][i]
        par = data["par_num"][i]
        line = data["line_num"][i]
        
        key = (block, par, line)
        
        if key not in lines_dict:
            lines_dict[key] = {
                "words": [],
                "left": data["left"][i],
                "top": data["top"][i],
                "height": data["height"][i],
                "width": 0,
            }
        
        word_data = {
            "text": txt,
            "left": data["left"][i],
            "top": data["top"][i],
            "width": data["width"][i],
            "height": data["height"][i],
            "conf": conf,
        }
        
        lines_dict[key]["words"].append(word_data)
        lines_dict[key]["left"] = min(lines_dict[key]["left"], data["left"][i])
        lines_dict[key]["top"] = min(lines_dict[key]["top"], data["top"][i])
        lines_dict[key]["height"] = max(lines_dict[key]["height"], data["height"][i])
    
    # Satrlarni birlashtirish va tartiblash
    lines = []
    all_heights = []
    
    for key, line_data in lines_dict.items():
        words = sorted(line_data["words"], key=lambda w: w["left"])
        
        # So'zlarni matunga birlashtirish
        text_parts = []
        prev_right = None
        
        for w in words:
            left = w["left"]
            
            if prev_right is not None:
                gap = left - prev_right
                # Agar katta bo'shliq bo'lsa, 2 ta space qo'shamiz
                if gap > line_data["height"] * 1.2:
                    text_parts.append("  ")
                else:
                    text_parts.append(" ")
            
            text_parts.append(w["text"])
            prev_right = left + w["width"]
        
        text = "".join(text_parts).strip()
        
        if text:
            lines.append({
                "text": text,
                "left": line_data["left"],
                "top": line_data["top"],
                "height": line_data["height"],
                "width": prev_right - line_data["left"] if prev_right else 0,
            })
            all_heights.append(line_data["height"])
    
    # Median balandlikni hisoblash (normal matn o'lchami uchun)
    median_height = sorted(all_heights)[len(all_heights) // 2] if all_heights else 16
    
    # Har bir satr uchun metadata qo'shamiz
    for line in lines:
        line["median_height"] = median_height
        line["height_ratio"] = line["height"] / median_height if median_height > 0 else 1.0
    
    # Top va left bo'yicha tartiblash
    lines.sort(key=lambda x: (x["top"], x["left"]))
    
    return lines


# ============================================================
# DOCX Builder with Formatting
# ============================================================

def _add_formatted_paragraph(
    doc: Document,
    line: Dict,
    min_left: int,
    prev_line: Optional[Dict] = None,
) -> None:
    """
    Formatted paragraph qo'shish:
    - Indent (left_indent)
    - Font size (height asosida)
    - Bold (sarlavhalar uchun)
    - List style (bullets va numbers)
    - Vertical spacing
    """
    text = line["text"].strip()
    if not text:
        return
    
    height = line["height"]
    median_h = line["median_height"]
    height_ratio = line["height_ratio"]
    
    # Indent hisoblash
    indent_px = max(0, line["left"] - min_left)
    left_indent = min(1.5, indent_px / 200.0)  # 200px ≈ 1 inch
    
    # Font size hisoblash
    font_size = _px_to_pt(height)
    
    # Sarlavha ekanligini aniqlash
    is_heading = _is_heading_text(text, height_ratio)
    
    # List item ekanligini aniqlash
    is_list, list_type = _is_bullet_or_numbered(text)
    
    # Paragraph yaratish
    if is_list:
        if list_type == "bullet":
            text = "• " + text.lstrip("•-·*○■► ").strip()
        p = doc.add_paragraph(style="List Bullet" if list_type == "bullet" else "List Number")
    elif is_heading:
        p = doc.add_paragraph(style="Heading 2" if height_ratio >= 1.5 else "Heading 3")
    else:
        p = doc.add_paragraph()
    
    # Vertical spacing (agar oldingi satrdan uzoq bo'lsa)
    if prev_line is not None:
        gap = line["top"] - (prev_line["top"] + prev_line["height"])
        if gap > median_h * 0.8:
            space_pt = min(12, max(3, gap / 5))
            p.paragraph_format.space_before = Pt(space_pt)
    
    # Indent
    if left_indent > 0 and not is_list:
        p.paragraph_format.left_indent = Inches(left_indent)
    
    # Text run
    run = p.add_run(text)
    _apply_font(run)
    
    # Font size
    if is_heading:
        run.bold = True
        run.font.size = Pt(min(18, max(12, font_size + 2)))
    else:
        run.font.size = Pt(min(14, max(10, font_size)))


# ============================================================
# Main DOCX Builder Functions
# ============================================================

def build_docx_bytes_from_image(
    path: Path,
    lang: str = "eng",
    fast_mode: bool = False,
) -> bytes:
    """
    Bitta rasmdan formatlangan DOCX yaratish
    """
    img = _safe_open_image(path)
    lines = _extract_lines_with_format(img, lang)
    
    doc = Document()
    
    if not lines:
        doc.add_paragraph("(OCR natijasi bo'sh)")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    
    # Min left - indent hisoblash uchun
    min_left = min(l["left"] for l in lines)
    
    prev_line = None
    for line in lines:
        _add_formatted_paragraph(doc, line, min_left, prev_line)
        prev_line = line
    
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx_bytes_from_images(
    paths: List[Path],
    lang: str = "eng",
    fast_mode: bool = False,
) -> bytes:
    """
    Ko'p rasmdan bitta formatlangan DOCX yaratish
    Har bir rasm - yangi sahifa
    """
    doc = Document()
    first = True
    
    for path in paths:
        if not first:
            doc.add_page_break()
        first = False
        
        img = _safe_open_image(path)
        lines = _extract_lines_with_format(img, lang)
        
        if not lines:
            doc.add_paragraph(f"(Sahifa {paths.index(path) + 1}: OCR natijasi bo'sh)")
            continue
        
        min_left = min(l["left"] for l in lines)
        
        prev_line = None
        for line in lines:
            _add_formatted_paragraph(doc, line, min_left, prev_line)
            prev_line = line
    
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# Text -> DOCX (simple, no formatting detection)
# ============================================================

def build_docx_bytes_from_text(text: str) -> bytes:
    """
    Oddiy matndan DOCX yaratish
    Bu funksiya hozircha oddiy qoladi, lekin keraksa yaxshilanishi mumkin
    """
    doc = Document()
    text = (text or "").replace("\r\n", "\n").strip()
    
    if not text:
        doc.add_paragraph("")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    
    lines = text.split("\n")
    blank = False
    
    for line in lines:
        line = line.rstrip()
        
        if not line.strip():
            blank = True
            continue
        
        if blank:
            doc.add_paragraph("")
            blank = False
        
        s = line.strip()
        
        # List yoki sarlavha aniqlash
        is_list, list_type = _is_bullet_or_numbered(s)
        is_heading = _is_heading_text(s)
        
        if is_list:
            if list_type == "bullet":
                s = "• " + s.lstrip("•-·*○■► ").strip()
            p = doc.add_paragraph(style="List Bullet" if list_type == "bullet" else "List Number")
        elif is_heading:
            p = doc.add_paragraph(style="Heading 2")
        else:
            p = doc.add_paragraph()
        
        run = p.add_run(s)
        _apply_font(run)
        
        if is_heading:
            run.bold = True
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(DEFAULT_FONT_SIZE)
    
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()